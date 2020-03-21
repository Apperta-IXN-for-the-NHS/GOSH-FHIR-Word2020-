# Patient Health Records Generator

# I developed a web application in Flask which uses a Python API for creating and
# updating Microsoft Word. I also used FHIR Parser which made it easier for me to
# access the information that I needed from HL7 FHIR.

# Sabina-Maria Mitroi - 18.03.2020

from docx import Document
from docx.shared import Inches
from fhir_parser import FHIR
from datetime import date


# Method that saves all the information for a patient in a document
def save_doc(text, dropdown):
    # Full list of patients
    fhir = FHIR()
    patients = fhir.get_all_patients()
    document = Document()
    new_line = " "

    # Find number of patients
    number_of_patients = 0
    for i in patients:
        number_of_patients = number_of_patients + 1

    # Index of Patient
    if dropdown == "Forename":
        result = forename(text)
    elif dropdown == "Surname":
        result = surname(text)
    else:
        result = unique_identifier(text)
    patient = patients[result]

    # Address of Patient
    address = patient.addresses[0]
    communications = patient.communications

    # Date Today
    today = date.today()

    # Marital Status
    marital_status_patient = patient.marital_status.marital_status

    # Age
    age = today.year - patient.birth_date.year - (
                (today.month, today.day) < (patient.birth_date.month, patient.birth_date.day))

    # Records of patients
    patient_data = {"Unique Identifier": patient.uuid.upper(), "Surname": patient.name.family,
                    "Forename": patient.name.given, "Gender": patient.gender.capitalize(),
                    "Birthday": str(patient.birth_date.day) + "." + str(patient.birth_date.month) + "." \
                                + str(patient.birth_date.year), "Age": str(age),
                    "Marital Status": str(marital_status_patient), "Address": address.full_address,
                    "Languages Spoken": communications.languages}

    # Records of a patient
    observations = fhir.get_patient_observations(patient_data["Unique Identifier"].lower())

    # Number of observations for 1 Patient
    count = 0
    for i in observations:
        count = count + 1

    # NHS Image
    LOGO = "1280px-NHS-Logo.svg.png"
    logo = document.add_picture(LOGO, width=Inches(1.50))
    document.add_heading(patient.name.full_name, 1)
    document.add_heading(new_line, 1)

    # Adding general information about a patient in a word document
    for data in patient_data:
        p = document.add_paragraph()
        p.add_run(data + ": ").bold = True
        p.add_run(patient_data[data])
    document.add_heading(new_line, 1)

    # Creating a table
    table = document.add_table(rows=1, cols=8)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Observation(s)       '
    hdr_cells[1].text = 'Number'
    hdr_cells[2].text = 'Type'
    hdr_cells[3].text = 'Status      '
    hdr_cells[5].text = 'Issued Date'
    hdr_cells[6].text = 'Code'
    hdr_cells[7].text = 'Result'

    # Adding data into the table
    for i in range(0, count):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i + 1)
        row_cells[1].text = observations[i].uuid.upper()
        row_cells[2].text = observations[i].type.capitalize()
        row_cells[3].text = observations[i].status.capitalize()
        row_cells[5].text = str(observations[i].issued_datetime)
        for component in observations[i].components:
            row_cells[6].text = component.code
            row_cells[7].text = component.display + ": " + component.quantity()

    # Save the document
    document.save(str(patient.name.given[0]) + "." + str(patient.name.family)
                  + "." + str(age) + '.docx')


# Method that finds a patient based on the forename
def forename(text):
    fhir = FHIR()
    patients = fhir.get_all_patients()
    number_of_patients = 0
    for i in patients:
        number_of_patients = number_of_patients + 1
    for index in range(0, number_of_patients):
        patient = patients[index]
        if text == patient.name.given:
            return index
    return -1


# Method that finds a patient based on the surname
def surname(text):
    fhir = FHIR()
    patients = fhir.get_all_patients()
    number_of_patients = 0
    for i in patients:
        number_of_patients = number_of_patients + 1
    for index in range(0, number_of_patients):
        patient = patients[index]
        if text == patient.name.family:
            return index
    return -1


# Method that finds a patient based on the unique identifier
def unique_identifier(text):
    fhir = FHIR()
    patients = fhir.get_all_patients()
    number_of_patients = 0
    for i in patients:
        number_of_patients = number_of_patients + 1
    for index in range(0, number_of_patients):
        patient = patients[index]
        if text == patient.uuid:
            return index
    return -1